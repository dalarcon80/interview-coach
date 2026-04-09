"""
Insights service for benchmark-driven CV strengthening.

This module is intentionally independent from Coach/Live. It benchmarks the
profile against curated leadership packs, keeps a durable workspace model for
Insights only, and produces approved context that can be persisted for future
use without wiring it into Brain in this phase.
"""

from __future__ import annotations

import base64
import copy
import json
import re
from dataclasses import dataclass
from hashlib import sha1
from io import BytesIO
from pathlib import Path
from typing import Any

from pipeline.steps.cv_analyzer import CVAnalyzer, CVAnalysisResult

_DATA_DIR = Path(__file__).with_name("insights_data")

_ACTION_TERMS = {
    "led",
    "built",
    "created",
    "designed",
    "founded",
    "implemented",
    "scaled",
    "drove",
    "owned",
    "modernized",
    "delivered",
    "expanded",
    "improved",
    "reduced",
    "increased",
    "accelerated",
    "defined",
    "guided",
    "standardized",
}

_LEADERSHIP_TERMS = {
    "team",
    "teams",
    "manager",
    "managers",
    "lead",
    "leadership",
    "director",
    "head",
    "reports",
    "regional",
    "global",
    "budget",
    "portfolio",
    "practice",
    "organization",
}

_TECHNICAL_TERMS = {
    "architecture",
    "architect",
    "platform",
    "pipeline",
    "cloud",
    "aws",
    "azure",
    "gcp",
    "spark",
    "dbt",
    "terraform",
    "lakehouse",
    "warehouse",
    "data",
    "ai",
    "ml",
    "genai",
    "modernization",
    "migration",
    "governance",
}

_ADVISORY_TERMS = {
    "client",
    "clients",
    "stakeholder",
    "stakeholders",
    "executive",
    "proposal",
    "portfolio",
    "accounts",
    "account",
    "consulting",
    "adoption",
    "advisory",
    "roadmap",
}

_SCOPE_TERMS = {
    "applications",
    "apps",
    "regions",
    "countries",
    "portfolio",
    "clients",
    "accounts",
    "budget",
    "direct",
    "indirect",
    "reports",
    "engineers",
    "architects",
    "teams",
    "program",
}


def _load_json(name: str) -> dict[str, Any]:
    return json.loads((_DATA_DIR / name).read_text(encoding="utf-8"))


def _s(value: Any) -> str:
    if value is None:
        return ""
    return str(value).strip()


def _i(value: Any) -> int:
    if isinstance(value, bool):
        return 0
    if isinstance(value, int):
        return value
    try:
        return int(str(value).strip())
    except Exception:
        return 0


def _list(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, list):
        raw = value
    else:
        raw = re.split(r"[\n,;]+", str(value))
    return [str(item).strip() for item in raw if str(item).strip()]


def _dedupe(values: list[str]) -> list[str]:
    seen: set[str] = set()
    ordered: list[str] = []
    for value in values:
        key = value.casefold()
        if not key or key in seen:
            continue
        seen.add(key)
        ordered.append(value)
    return ordered


def _truncate(text: str, limit: int) -> str:
    clean = _s(text)
    if len(clean) <= limit:
        return clean
    return clean[: limit - 3].rstrip() + "..."


def _tokens(text: str) -> set[str]:
    return {
        token
        for token in re.findall(r"[a-zA-Z][a-zA-Z0-9+#.&/-]{2,}", _s(text).lower())
        if len(token) > 2
    }


def _contains_metric(text: str) -> bool:
    return bool(
        re.search(
            r"(\d+%|\$\d+|\d+\+|\d+\s*(?:m|k|million|billion)|\d+\s*(?:people|apps|applications|clients|accounts|projects|regions|countries|reports|teams))",
            text,
            re.I,
        )
    )


def _contains_scope(text: str) -> bool:
    lower = text.lower()
    return any(term in lower for term in _SCOPE_TERMS) or bool(
        re.search(r"\b\d+\s*(?:direct|indirect|regional|global|client|account|application|team)\b", lower)
    )


def _contains_action(text: str) -> bool:
    return any(token in _ACTION_TERMS for token in _tokens(text))


def _contains_ownership(text: str) -> bool:
    lower = text.lower()
    return any(term in lower for term in ("owned", "led", "managed", "guided", "defined", "founded", "built"))


def _normalize_lines(text: str) -> list[str]:
    lines = re.split(r"[\n\r]+|(?:\s*[-•]\s+)", text)
    normalized: list[str] = []
    for line in lines:
        clean = " ".join(str(line or "").split()).strip(" -•\t")
        if len(clean) >= 18:
            normalized.append(clean)
    return normalized


def _headline(text: str) -> str:
    return _truncate(_s(text).split(".")[0], 90) or "Candidate"


def _fingerprint(payload: dict[str, Any]) -> str:
    encoded = json.dumps(payload, sort_keys=True, ensure_ascii=True, default=str)
    return sha1(encoded.encode("utf-8")).hexdigest()


def _slug(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", "_", _s(text).lower()).strip("_")


def _sanitize_filename(value: str) -> str:
    safe = re.sub(r"[^a-zA-Z0-9._-]+", "-", value.strip())
    return safe.strip("-_.") or "candidate"


def _delta_template(dimension: str, severity: str = "medium") -> dict[str, int]:
    base = {
        "global": 3,
        "roleFit": 2,
        "proofStrength": 2,
        "cvRepresentationQuality": 1,
    }
    if dimension == "impact_outcomes":
        base = {"global": 5, "roleFit": 2, "proofStrength": 6, "cvRepresentationQuality": 1}
    elif dimension == "scope_complexity":
        base = {"global": 4, "roleFit": 4, "proofStrength": 3, "cvRepresentationQuality": 1}
    elif dimension == "functional_technical_depth":
        base = {"global": 4, "roleFit": 5, "proofStrength": 4, "cvRepresentationQuality": 1}
    elif dimension == "role_positioning":
        base = {"global": 3, "roleFit": 5, "proofStrength": 1, "cvRepresentationQuality": 4}
    elif dimension == "cv_clarity_ats":
        base = {"global": 2, "roleFit": 1, "proofStrength": 1, "cvRepresentationQuality": 5}
    elif dimension == "influence_collaboration":
        base = {"global": 3, "roleFit": 4, "proofStrength": 2, "cvRepresentationQuality": 1}
    elif dimension == "evidence_quality":
        base = {"global": 4, "roleFit": 2, "proofStrength": 5, "cvRepresentationQuality": 2}

    multiplier = 1.0
    if severity == "high":
        multiplier = 1.35
    elif severity == "low":
        multiplier = 0.7
    return {key: max(1, round(value * multiplier)) for key, value in base.items()}


def _answer_schema(question_family: str) -> dict[str, Any]:
    if question_family == "project_recovery":
        return {
            "fields": [
                "project_name",
                "business_problem",
                "architecture_or_system",
                "scale",
                "key_tradeoffs",
                "your_ownership",
                "outcome",
            ],
            "format_hint": "Project + architecture + scale + trade-offs + your ownership + outcome",
        }
    if question_family == "impact_recovery":
        return {
            "fields": ["baseline", "metric", "achieved_change", "timeframe", "your_role"],
            "format_hint": "Baseline + metric + change achieved + timeframe + your role",
        }
    if question_family == "scope_recovery":
        return {
            "fields": ["team_or_program", "scope", "regions_or_clients", "your_ownership", "outcome"],
            "format_hint": "Team/program + scale + scope + ownership + result",
        }
    if question_family in {"target_positioning_clarification", "role_positioning"}:
        return {
            "fields": ["principal_track", "strongest_domain", "scope", "why_principal", "preferred_positioning_sentence"],
            "format_hint": "Track + strongest domain + scope + why principal + positioning sentence",
        }
    if question_family == "managerial":
        return {
            "fields": ["scope", "team_size", "regions_or_portfolio", "decision", "outcome"],
            "format_hint": "Scope + decision + measurable result",
        }
    if question_family in {"technical", "architecture_depth_recovery"}:
        return {
            "fields": ["system_or_program", "architecture", "tradeoffs", "scale", "outcome"],
            "format_hint": "System + architecture + tradeoff + scale + result",
        }
    if question_family in {"consulting_advisory", "business_alignment_recovery"}:
        return {
            "fields": ["client_or_stakeholder", "problem", "intervention", "business_outcome"],
            "format_hint": "Stakeholder problem + intervention + business outcome",
        }
    return {
        "fields": ["result", "metric", "timeframe", "ownership"],
        "format_hint": "Result + metric + timeframe + your ownership",
    }


def _answer_guidance(question_family: str) -> str:
    if question_family == "project_recovery":
        return "Focus on one principal-level project. Describe the business problem, the system or architecture, the key trade-offs, your ownership, and the measurable outcome."
    if question_family == "impact_recovery":
        return "Be explicit about the before state, the metric you moved, the magnitude of change, the timeframe, and what part of that result was yours."
    if question_family == "scope_recovery":
        return "Name the team, program, or platform scope clearly and make the scale concrete with regions, clients, applications, or users."
    if question_family in {"target_positioning_clarification", "role_positioning"}:
        return "Write the positioning the way you would want a recruiter or hiring manager to understand it in one pass."
    if question_family == "managerial":
        return "Include team or portfolio size, your ownership, the decision you drove, and the concrete result."
    if question_family in {"technical", "architecture_depth_recovery"}:
        return "Describe the system or program, the technical choices, the scale involved, and what improved because of your work."
    if question_family in {"consulting_advisory", "business_alignment_recovery"}:
        return "Anchor the answer in one client or stakeholder problem, what you changed, and the business result you enabled."
    return "Be explicit about the metric, the timeframe, and the part of the result that was yours."


def _answer_example(question_family: str) -> str:
    if question_family == "project_recovery":
        return "Project: Core banking data platform modernization. Business problem: fragmented batch processing slowed decision-making. Architecture/system: I defined the target lakehouse pattern on Azure with reusable ingestion and semantic layers. Scale: 120+ pipelines across 6 business units. Key trade-offs: we chose phased domain migration over big-bang cutover to protect critical reporting. Ownership: I owned the technical direction and decision reviews. Outcome: latency dropped from 8 hours to 40 minutes and platform run cost fell 32%."
    if question_family == "impact_recovery":
        return "Baseline: provisioning a new data product took 6 weeks. Metric: time-to-production. Achieved change: reduced to 8 days. Timeframe: within two quarters. Your role: I defined the operating model, architecture guardrails, and rollout sequence."
    if question_family == "scope_recovery":
        return "Team/program: regional data modernization portfolio. Scope: 4 platform squads and 18 critical applications. Regions/clients: 3 countries and 5 enterprise clients. Ownership: I led the technical roadmap and cross-team decision cadence. Outcome: delivery predictability improved 25%."
    if question_family in {"target_positioning_clarification", "role_positioning"}:
        return "Principal track: principal platform architect. Strongest domain: enterprise data modernization. Scope: multi-region platform transformation for regulated environments. Why principal: I operate through architecture direction, technical influence, and reusable platform patterns rather than line management. Preferred positioning sentence: Principal-level data platform leader focused on modernization, architecture depth, and cross-team technical influence."
    if question_family == "managerial":
        return "I led 4 regional managers and 36 engineers across 3 countries, reset the operating cadence, and improved delivery predictability by 28% in two quarters."
    if question_family in {"technical", "architecture_depth_recovery"}:
        return "I owned the lakehouse migration target state on Azure, chose dbt plus Delta patterns over direct lift-and-shift, supported 120+ pipelines, and cut data latency from 9 hours to 45 minutes."
    if question_family in {"consulting_advisory", "business_alignment_recovery"}:
        return "I reframed the client's modernization roadmap around revenue-critical journeys, aligned executive sponsors, and unlocked a phased program that reduced onboarding time by 35%."
    return "I standardized delivery governance for our data programs and reduced time-to-impact by 30% within 12 months."


def _dimension_missing_defaults(dimension_id: str) -> list[str]:
    defaults = {
        "impact_outcomes": [
            "Missing quantified outcome",
            "Missing baseline vs result",
            "Missing attribution to your work",
        ],
        "scope_complexity": [
            "Missing team, program, or platform scale",
            "Missing regional, client, or application scope",
            "Missing complexity markers",
        ],
        "functional_technical_depth": [
            "Missing architecture choice",
            "Missing trade-offs",
            "Missing system or platform detail",
        ],
        "role_positioning": [
            "Missing explicit principal framing",
            "Missing target-alignment sentence",
            "Missing strongest domain positioning",
        ],
        "evidence_quality": [
            "Missing clear ownership",
            "Missing concrete proof point",
            "Missing sharper example",
        ],
        "influence_collaboration": [
            "Missing cross-team influence example",
            "Missing stakeholder impact",
        ],
        "cv_clarity_ats": [
            "Missing sharper summary structure",
            "Missing clearer evidence packaging",
        ],
    }
    return defaults.get(dimension_id, ["Missing stronger evidence"])


def _dimension_next_best_action(dimension_id: str) -> str:
    actions = {
        "impact_outcomes": "Answer an impact question with baseline, metric, achieved change, timeframe, and your ownership.",
        "scope_complexity": "Add one scope example with team, platform, client, region, or application scale.",
        "functional_technical_depth": "Answer one project or architecture-depth question with system detail and trade-offs.",
        "role_positioning": "Clarify the exact principal positioning sentence you want the benchmark and CV to optimize for.",
        "evidence_quality": "Add one concrete proof point with outcome, ownership, and timeframe.",
        "influence_collaboration": "Add one example of technical or stakeholder influence that changed direction or execution.",
        "cv_clarity_ats": "Approve one rewrite and regenerate the role variant so the strongest evidence is easier to scan.",
    }
    return actions.get(dimension_id, "Answer the next guided question to strengthen this area.")


def _source_type(source: str) -> str:
    if source == "answer":
        return "user_answer"
    if source in {"candidate_summary", "achievement", "skill"}:
        return "imported_profile"
    if source.startswith("cv"):
        return "cv"
    return "system_extraction"


def _confidence_from_strength(strength: str) -> str:
    if strength == "strong":
        return "high"
    if strength == "moderate":
        return "medium"
    return "low"


def _normalize_candidate(profile: dict[str, Any] | None) -> dict[str, Any]:
    source = profile or {}
    return {
        "name": _s(source.get("name")),
        "current_role": _s(source.get("current_role") or source.get("currentRole")),
        "years_experience": _i(source.get("years_experience") or source.get("yearsExperience")),
        "skills": _list(source.get("skills")),
        "education": _s(source.get("education")),
        "languages": _list(source.get("languages")),
        "certifications": _list(source.get("certifications")),
        "summary": _s(source.get("summary")),
        "achievements": _list(source.get("achievements")),
        "target_role": _s(source.get("target_role") or source.get("targetRole")),
        "industry": _s(source.get("industry")),
        "location": _s(source.get("location")),
        "cv_text": _s(source.get("cv_text") or source.get("cvText")),
        "profile_id": _s(source.get("profile_id")) or None,
        "insights_context_summary": _s(source.get("insights_context_summary")),
        "insights_focus_areas": _list(source.get("insights_focus_areas")),
        "insights_reusable_evidence": _list(source.get("insights_reusable_evidence")),
    }


def _normalize_company(info: dict[str, Any] | None) -> dict[str, Any]:
    source = info or {}
    return {
        "name": _s(source.get("name") or source.get("companyName")),
        "industry": _s(source.get("industry")),
        "role_title": _s(source.get("role_title") or source.get("roleTitle") or source.get("positionTitle")),
        "role_level": _s(source.get("role_level") or source.get("roleLevel")),
        "role_requirements": _list(
            source.get("role_requirements") or source.get("roleRequirements") or source.get("positionRequirements")
        ),
        "role_responsibilities": _list(source.get("role_responsibilities") or source.get("roleResponsibilities")),
        "interview_focus": _list(source.get("interview_focus") or source.get("interviewFocus")),
        "job_description": _s(source.get("job_description") or source.get("jobDescription") or source.get("positionDescription")),
        "culture": _s(source.get("culture") or source.get("companyCulture")),
    }


def _normalize_interviewer(info: dict[str, Any] | None) -> dict[str, Any]:
    source = info or {}
    return {
        "name": _s(source.get("name")),
        "role_title": _s(source.get("role_title") or source.get("roleTitle")),
        "company": _s(source.get("company")),
        "expertise": _list(source.get("expertise")),
        "likely_focus_areas": _list(source.get("likely_focus_areas") or source.get("likelyFocusAreas")),
        "notes": _s(source.get("notes")),
    }


@dataclass
class BenchmarkResolution:
    benchmark_source: dict[str, Any]
    merged_pack: dict[str, Any]
    support_level: str


class InsightsService:
    def __init__(self) -> None:
        self.cv_analyzer = CVAnalyzer.from_environment()
        self.global_rubric = _load_json("global_rubric.json")
        self.archetype_packs = _load_json("archetype_packs.json")
        self.role_family_packs = _load_json("role_family_packs.json")
        self.seniority_packs = _load_json("seniority_packs.json")

    async def analyze(
        self,
        candidate_profile: dict[str, Any] | None,
        company_info: dict[str, Any] | None,
        interviewer_profile: dict[str, Any] | None,
        cv_text: str,
        language: str = "en",
        answers: dict[str, str] | None = None,
        target_role_override: str | None = None,
        archetype_override: str | None = None,
        seniority_override: str | None = None,
        specialty_ids: list[str] | None = None,
    ) -> dict[str, Any]:
        candidate = _normalize_candidate(candidate_profile)
        company = _normalize_company(company_info)
        interviewer = _normalize_interviewer(interviewer_profile)
        candidate["cv_text"] = _s(cv_text) or candidate["cv_text"]
        answers = { _s(k): _s(v) for k, v in (answers or {}).items() if _s(k) and _s(v) }
        specialty_ids = [item for item in (specialty_ids or []) if _s(item)]

        cv_result: CVAnalysisResult | None = None
        if len(candidate["cv_text"]) >= 50:
            cv_result = await self.cv_analyzer.analyze(candidate["cv_text"])

        candidate = self._enrich_candidate_from_cv(candidate, cv_result)
        resolution = self._resolve_benchmark(
            candidate=candidate,
            company=company,
            answers=answers,
            target_role_override=target_role_override,
            archetype_override=archetype_override,
            seniority_override=seniority_override,
            specialty_ids=specialty_ids,
        )
        canonical = self._build_canonical(candidate, company, answers, cv_result)
        evidence_cards = self._build_evidence_cards(
            canonical=canonical,
            merged_pack=resolution.merged_pack,
            support_level=resolution.support_level,
        )
        signal_snapshot = self._evaluate_signals(evidence_cards, resolution.merged_pack)
        dimension_states = self._score_dimensions(
            candidate=candidate,
            canonical=canonical,
            signal_snapshot=signal_snapshot,
            merged_pack=resolution.merged_pack,
            evidence_cards=evidence_cards,
        )
        primary_scores = self._score_primary_scores(
            candidate=candidate,
            canonical=canonical,
            signal_snapshot=signal_snapshot,
            dimension_states=dimension_states,
            support_level=resolution.support_level,
        )
        gap_map = self._build_gap_map(signal_snapshot, dimension_states)
        questions = self._build_questions(
            signal_snapshot=signal_snapshot,
            gap_map=gap_map,
            benchmark_source=resolution.benchmark_source,
            primary_scores=primary_scores,
            answers=answers,
        )
        recommended_profile = self._build_recommended_profile(candidate, company, answers, evidence_cards, primary_scores)
        cv_variants = self._build_cv_variants(
            candidate=candidate,
            company=company,
            benchmark_source=resolution.benchmark_source,
            recommended_profile=recommended_profile,
            evidence_cards=evidence_cards,
            gap_map=gap_map,
        )
        proposed_changes = self._build_proposed_changes(candidate, recommended_profile)
        next_actions, improvement_plan = self._build_improvement_plan(
            benchmark_source=resolution.benchmark_source,
            primary_scores=primary_scores,
            gap_map=gap_map,
            questions=questions,
            proposed_changes=proposed_changes,
            cv_variants=cv_variants,
        )
        top_strengths = self._build_top_strengths(signal_snapshot, evidence_cards)
        top_gaps = self._build_top_gaps(gap_map)
        interpretation = self._build_interpretation(
            benchmark_source=resolution.benchmark_source,
            primary_scores=primary_scores,
            top_strengths=top_strengths,
            top_gaps=top_gaps,
            support_level=resolution.support_level,
        )
        approved_context_preview = self._build_approved_context_preview(
            benchmark_source=resolution.benchmark_source,
            primary_scores=primary_scores,
            evidence_cards=evidence_cards,
            recommended_profile=recommended_profile,
            selected_evidence_ids=[],
        )

        workspace_state = "active"
        return {
            "mode": cv_result.mode if cv_result else "fallback",
            "analysis_summary": approved_context_preview["summary"],
            "benchmark_source": resolution.benchmark_source,
            "support_level": resolution.support_level,
            "workspace_state": workspace_state,
            "primary_scores": primary_scores["primary_scores"],
            "global_score": primary_scores["overall_match"],
            "overall_match": primary_scores["overall_match"],
            "coverage_pct": primary_scores["coverage_pct"],
            "confidence": primary_scores["confidence"],
            "score_delta_available": sum(step["estimated_delta"].get("global", 0) for step in next_actions),
            "top_strengths": top_strengths,
            "top_gaps": top_gaps,
            "interpretation": interpretation,
            "next_actions": next_actions,
            "improvement_plan": improvement_plan,
            "dimension_states": dimension_states,
            "required_signals": signal_snapshot["required_signals"],
            "supporting_signals": signal_snapshot["supporting_signals"],
            "differentiator_signals": signal_snapshot["differentiator_signals"],
            "anti_signals": signal_snapshot["anti_signals"],
            "not_applicable_signals": signal_snapshot["not_applicable_signals"],
            "gap_map": gap_map,
            "evidence_cards": evidence_cards,
            "questions": questions,
            "proposed_changes": proposed_changes,
            "recommended_profile": recommended_profile,
            "approved_context_preview": approved_context_preview,
            "insights_context_summary": approved_context_preview["summary"],
            "cv_health": self._build_cv_health(dimension_states),
            "role_match_summary": self._build_role_match_summary(resolution.benchmark_source, primary_scores),
            "cv_variants": cv_variants,
            "answers": answers,
            "input_snapshot": {
                "candidate_profile": candidate,
                "company_info": company,
                "interviewer_profile": interviewer,
                "cv_text": candidate["cv_text"],
                "language": language,
                "answers": answers,
                "target_role_override": target_role_override or "",
                "archetype_override": archetype_override or "",
                "seniority_override": seniority_override or "",
                "specialty_ids": specialty_ids,
            },
            "signal_snapshot": signal_snapshot,
            "score_history": [],
        }

    def _enrich_candidate_from_cv(self, candidate: dict[str, Any], cv_result: CVAnalysisResult | None) -> dict[str, Any]:
        enriched = copy.deepcopy(candidate)
        if not cv_result or not cv_result.success:
            enriched["skills"] = _dedupe(enriched["skills"])
            enriched["achievements"] = _dedupe(enriched["achievements"])
            return enriched

        profile = cv_result.profile
        if not enriched["current_role"]:
            enriched["current_role"] = profile.current_role
        if not enriched["years_experience"]:
            enriched["years_experience"] = profile.years_experience
        if not enriched["summary"]:
            enriched["summary"] = profile.summary
        enriched["skills"] = _dedupe(enriched["skills"] + profile.skills + profile.technical_stack)
        enriched["achievements"] = _dedupe(enriched["achievements"] + profile.achievements + profile.metrics)
        return enriched

    def _resolve_benchmark(
        self,
        *,
        candidate: dict[str, Any],
        company: dict[str, Any],
        answers: dict[str, str],
        target_role_override: str | None,
        archetype_override: str | None,
        seniority_override: str | None,
        specialty_ids: list[str],
    ) -> BenchmarkResolution:
        target_role = (
            _s(target_role_override)
            or answers.get("target_role")
            or candidate.get("target_role")
            or company.get("role_title")
            or candidate.get("current_role")
        )
        normalized_target_role = _slug(target_role)

        family_pack = self._resolve_role_family(
            target_role,
            candidate=candidate,
            company=company,
            answers=answers,
        )
        support_level = family_pack.get("support_level", "unsupported") if family_pack else "unsupported"
        if not family_pack:
            family_pack = {
                "id": "generic_unsupported",
                "label": target_role or "Unsupported role",
                "required_signals": [],
                "supporting_signals": [],
                "differentiator_signals": [],
                "anti_signals": [],
                "not_applicable_defaults": [],
            }

        archetype_pack = self._resolve_archetype(target_role, family_pack, archetype_override)
        seniority_pack = self._resolve_seniority(target_role, seniority_override)

        merged_dimensions = []
        total_weight = 0.0
        for base_dimension in self.global_rubric["dimensions"]:
            dimension = copy.deepcopy(base_dimension)
            dimension["weight"] = float(base_dimension["weight"])
            for pack in (archetype_pack, family_pack, seniority_pack):
                override = (pack.get("dimension_overrides") or {}).get(dimension["id"], {})
                if _s(override.get("label")):
                    dimension["label"] = override["label"]
                if _s(override.get("description")):
                    dimension["description"] = override["description"]
                dimension["weight"] += float(override.get("weight_delta") or 0)
            dimension["weight"] = max(0.01, dimension["weight"])
            total_weight += dimension["weight"]
            merged_dimensions.append(dimension)
        for dimension in merged_dimensions:
            dimension["weight"] = round(dimension["weight"] / total_weight, 4)

        merged_pack = {
            "dimensions": merged_dimensions,
            "required_signals": copy.deepcopy(family_pack.get("required_signals") or []),
            "supporting_signals": copy.deepcopy(family_pack.get("supporting_signals") or []),
            "differentiator_signals": copy.deepcopy(family_pack.get("differentiator_signals") or []),
            "anti_signals": copy.deepcopy(family_pack.get("anti_signals") or []),
            "not_applicable_defaults": copy.deepcopy(family_pack.get("not_applicable_defaults") or []),
            "question_family_order": copy.deepcopy(archetype_pack.get("question_family_order") or ["universal"]),
        }

        fingerprint = _fingerprint(
            {
                "target_role": target_role,
                "family": family_pack["id"],
                "archetype": archetype_pack["id"],
                "seniority": seniority_pack["id"],
                "specialty_ids": specialty_ids,
                "global_rubric_version": self.global_rubric["version"],
                "archetype_pack_version": self.archetype_packs["version"],
                "role_family_pack_version": self.role_family_packs["version"],
                "seniority_pack_version": self.seniority_packs["version"],
                "resolver_version": self.global_rubric["resolver_version"],
            }
        )

        benchmark_source = {
            "target_role": target_role,
            "normalized_target_role": normalized_target_role,
            "family": family_pack["label"],
            "family_pack_id": family_pack["id"],
            "archetype": archetype_pack["label"],
            "archetype_pack_id": archetype_pack["id"],
            "seniority": seniority_pack["label"],
            "seniority_pack_id": seniority_pack["id"],
            "specialty_ids": specialty_ids,
            "support_level": support_level,
            "versions": {
                "global_rubric_version": self.global_rubric["version"],
                "archetype_pack_version": self.archetype_packs["version"],
                "role_family_pack_version": self.role_family_packs["version"],
                "seniority_pack_version": self.seniority_packs["version"],
                "resolver_version": self.global_rubric["resolver_version"],
            },
            "benchmark_source_fingerprint": fingerprint,
        }

        return BenchmarkResolution(
            benchmark_source=benchmark_source,
            merged_pack=merged_pack,
            support_level=support_level,
        )

    def _resolve_role_family(
        self,
        target_role: str,
        *,
        candidate: dict[str, Any],
        company: dict[str, Any],
        answers: dict[str, str],
    ) -> dict[str, Any] | None:
        lower = _s(target_role).lower()
        if not lower:
            lower = _s(
                answers.get("target_role")
                or candidate.get("target_role")
                or company.get("role_title")
                or candidate.get("current_role")
            ).lower()
            if not lower:
                return None

        best_pack: dict[str, Any] | None = None
        best_score = 0
        for pack in self.role_family_packs["packs"]:
            aliases = " ".join(pack.get("aliases", []))
            pack_text = f"{pack['label']} {aliases}".lower()
            score = 0
            if lower == pack["label"].lower() or lower in [alias.lower() for alias in pack.get("aliases", [])]:
                score += 20
            role_tokens = _tokens(lower)
            pack_tokens = _tokens(pack_text)
            score += len(role_tokens & pack_tokens) * 4
            if score > best_score:
                best_score = score
                best_pack = pack
        if "principal" in lower or "staff" in lower:
            evidence_text = " ".join(
                part
                for part in [
                    target_role,
                    candidate.get("current_role", ""),
                    candidate.get("target_role", ""),
                    candidate.get("summary", ""),
                    candidate.get("cv_text", ""),
                    " ".join(candidate.get("skills", [])),
                    " ".join(candidate.get("achievements", [])),
                    company.get("role_title", ""),
                    company.get("job_description", ""),
                    " ".join(company.get("role_requirements", [])),
                ]
                if part
            ).lower()

            candidate_scores = {
                "principal_data_engineering": 0,
                "principal_architecture": 0,
                "principal_platform_modernization": 0,
            }
            if any(term in evidence_text for term in ("data", "pipeline", "lakehouse", "warehouse", "dbt", "spark", "data platform")):
                candidate_scores["principal_data_engineering"] += 5
            if any(term in evidence_text for term in ("architect", "architecture", "target state", "reference architecture", "governance")):
                candidate_scores["principal_architecture"] += 5
            if any(term in evidence_text for term in ("modernization", "migration", "platform modernization", "transformation", "cloud migration", "platform")):
                candidate_scores["principal_platform_modernization"] += 5

            for pack in self.role_family_packs["packs"]:
                if not pack["id"].startswith("principal_") or pack["id"] == "principal_generic_technical_leadership":
                    continue
                score = candidate_scores.get(pack["id"], 0)
                keywords = [
                    keyword
                    for bucket in ("required_signals", "supporting_signals", "differentiator_signals")
                    for signal in pack.get(bucket, [])
                    for keyword in signal.get("keywords", [])
                ]
                score += sum(1 for keyword in keywords if keyword.lower() in evidence_text)
                candidate_scores[pack["id"]] = score

            best_principal_family, best_principal_score = max(candidate_scores.items(), key=lambda item: item[1])
            if best_principal_score >= 4:
                return copy.deepcopy(
                    next(pack for pack in self.role_family_packs["packs"] if pack["id"] == best_principal_family)
                )
            return copy.deepcopy(
                next(
                    pack
                    for pack in self.role_family_packs["packs"]
                    if pack["id"] == "principal_generic_technical_leadership"
                )
            )
        if best_score >= 8:
            return copy.deepcopy(best_pack)
        if "director" in lower and ("data" in lower or "ingenier" in lower or "engineering" in lower):
            derived = copy.deepcopy(self.role_family_packs["packs"][1])
            derived["support_level"] = "derived"
            return derived
        return None

    def _resolve_archetype(self, target_role: str, family_pack: dict[str, Any], override: str | None) -> dict[str, Any]:
        if _s(override):
            for pack in self.archetype_packs["packs"]:
                if pack["id"] == override:
                    return copy.deepcopy(pack)
        preferred = _s(family_pack.get("preferred_archetype"))
        if preferred:
            match = next((pack for pack in self.archetype_packs["packs"] if pack["id"] == preferred), None)
            if match:
                return copy.deepcopy(match)
        lower = _s(target_role).lower()
        if family_pack["id"] == "data_ai_consulting":
            return copy.deepcopy(next(pack for pack in self.archetype_packs["packs"] if pack["id"] == "consulting_advisory"))
        if "principal" in lower or "staff" in lower:
            return copy.deepcopy(next(pack for pack in self.archetype_packs["packs"] if pack["id"] == "technical_leadership_principal"))
        if "manager" in lower and "director" not in lower and "head" not in lower:
            return copy.deepcopy(next(pack for pack in self.archetype_packs["packs"] if pack["id"] == "people_manager"))
        return copy.deepcopy(next(pack for pack in self.archetype_packs["packs"] if pack["id"] == "director_head"))

    def _resolve_seniority(self, target_role: str, override: str | None) -> dict[str, Any]:
        if _s(override):
            for pack in self.seniority_packs["packs"]:
                if pack["id"] == override:
                    return copy.deepcopy(pack)
        lower = _s(target_role).lower()
        for pack in self.seniority_packs["packs"]:
            if any(alias in lower for alias in pack.get("aliases", [])):
                return copy.deepcopy(pack)
        if "principal" in lower or "staff" in lower:
            return copy.deepcopy(next(pack for pack in self.seniority_packs["packs"] if pack["id"] == "principal"))
        return copy.deepcopy(next(pack for pack in self.seniority_packs["packs"] if pack["id"] == "director"))

    def _build_canonical(
        self,
        candidate: dict[str, Any],
        company: dict[str, Any],
        answers: dict[str, str],
        cv_result: CVAnalysisResult | None,
    ) -> dict[str, Any]:
        source_lines: list[dict[str, Any]] = []

        def add_lines(kind: str, values: list[str]) -> None:
            for value in values:
                clean = " ".join(value.split()).strip()
                if len(clean) >= 18:
                    source_lines.append({"source": kind, "text": clean})

        add_lines("candidate_summary", [candidate.get("summary", "")])
        add_lines("achievement", candidate.get("achievements", []))
        add_lines("skill", candidate.get("skills", []))
        add_lines("answer", list(answers.values()))
        add_lines("cv", _normalize_lines(candidate.get("cv_text", ""))[:30])
        if cv_result and cv_result.success:
            add_lines("cv_metric", cv_result.profile.metrics[:10])
            add_lines("cv_stack", cv_result.profile.technical_stack[:10])
            add_lines("cv_leadership", cv_result.profile.leadership_roles[:8])

        return {
            "target_role": answers.get("target_role") or candidate.get("target_role") or company.get("role_title") or candidate.get("current_role"),
            "summary": candidate.get("summary", ""),
            "skills": _dedupe(candidate.get("skills", [])),
            "achievements": _dedupe(candidate.get("achievements", [])),
            "cv_text": candidate.get("cv_text", ""),
            "source_lines": source_lines,
            "metrics": [line["text"] for line in source_lines if _contains_metric(line["text"])],
        }

    def _build_evidence_cards(
        self,
        *,
        canonical: dict[str, Any],
        merged_pack: dict[str, Any],
        support_level: str,
    ) -> list[dict[str, Any]]:
        cards: list[dict[str, Any]] = []
        seen: set[str] = set()
        for index, item in enumerate(canonical["source_lines"]):
            text = item["text"]
            key = text.casefold()
            if key in seen:
                continue
            seen.add(key)
            evidence_type = self._evidence_type(text, item["source"])
            proof = {
                "metrics_present": _contains_metric(text),
                "scope_present": _contains_scope(text),
                "ownership_present": _contains_ownership(text),
                "recency_present": bool(re.search(r"\b(20\d{2}|12 months|months|years)\b", text.lower())),
            }
            state = "inferred"
            if item["source"] == "answer":
                state = "needs_confirmation"
            strength = "strong" if sum(1 for value in proof.values() if value) >= 3 else "moderate" if sum(1 for value in proof.values() if value) >= 2 else "weak"
            signal_ids = []
            dimensions = set()
            for bucket in ("required_signals", "supporting_signals", "differentiator_signals"):
                for signal in merged_pack.get(bucket, []):
                    if _tokens(text) & set(signal.get("keywords", [])):
                        signal_ids.append(signal["id"])
                        dimensions.add(signal["dimension"])

            cards.append(
                {
                    "id": f"evidence::{index}",
                    "type": evidence_type,
                    "state": state,
                    "source": _source_type(item["source"]),
                    "summary": _truncate(text, 160),
                    "raw_evidence": text,
                    "dimensions": sorted(dimensions),
                    "signal_ids": signal_ids,
                    "role_relevance": {
                        "archetype": [],
                        "family": [],
                        "seniority": [],
                    },
                    "proof": proof,
                    "strength": strength,
                    "confidence": _confidence_from_strength(strength),
                    "approval_status": "needs_follow_up" if state == "needs_confirmation" else "draft",
                    "estimated_delta": _delta_template(sorted(dimensions)[0] if dimensions else "evidence_quality", "medium"),
                    "support_level": support_level,
                }
            )
        return cards

    def _evidence_type(self, text: str, source: str) -> str:
        tokens = _tokens(text)
        if source == "cv" and len(text.split()) < 12:
            return "cv_quality_evidence"
        if tokens & _ADVISORY_TERMS:
            return "advisory_evidence"
        if tokens & _TECHNICAL_TERMS:
            return "architecture_evidence"
        if tokens & _LEADERSHIP_TERMS:
            return "leadership_evidence"
        if _contains_action(text) and _contains_metric(text):
            return "impact_evidence"
        if _contains_action(text):
            return "project_evidence"
        return "delivery_evidence"

    def _evaluate_signals(self, evidence_cards: list[dict[str, Any]], merged_pack: dict[str, Any]) -> dict[str, Any]:
        state_weights = {
            "draft": 0.15,
            "inferred": 0.35,
            "needs_confirmation": 0.55,
            "approved": 1.0,
            "rejected": 0.0,
            "superseded": 0.0,
            "indexed": 1.0,
        }

        def build_signal_results(signals: list[dict[str, Any]], tier: str) -> list[dict[str, Any]]:
            results = []
            for signal in signals:
                matched = [card for card in evidence_cards if signal["id"] in card["signal_ids"]]
                coverage = min(100, round(sum(state_weights.get(card["state"], 0) * 55 for card in matched)))
                status = "covered" if coverage >= 70 else "partial" if coverage >= 35 else "missing"
                confidence_score = min(100, round(sum(state_weights.get(card["state"], 0) * 50 for card in matched)))
                results.append(
                    {
                        "id": signal["id"],
                        "label": signal["label"],
                        "dimension": signal["dimension"],
                        "tier": tier,
                        "status": status,
                        "coverage": coverage,
                        "confidence": confidence_score,
                        "matched_evidence_ids": [card["id"] for card in matched[:4]],
                        "question_family": signal.get("question_family", "universal"),
                        "question_template": signal.get("question_template", ""),
                        "expected_evidence": signal.get("expected_evidence", ""),
                    }
                )
            return results

        required = build_signal_results(merged_pack.get("required_signals", []), "required")
        supporting = build_signal_results(merged_pack.get("supporting_signals", []), "supporting")
        differentiators = build_signal_results(merged_pack.get("differentiator_signals", []), "differentiator")
        anti_signals = []
        for signal in merged_pack.get("anti_signals", []):
            matched = [card for card in evidence_cards if _tokens(card["raw_evidence"]) & set(signal.get("keywords", []))]
            active = False
            rule = signal.get("rule", "")
            if rule == "leadership_without_scope":
                active = any(card["type"] == "leadership_evidence" and not card["proof"]["scope_present"] for card in matched)
            elif rule == "modernization_without_technical_detail":
                active = any(card["type"] == "project_evidence" and not card["proof"]["ownership_present"] for card in matched)
            elif rule == "architecture_without_outcome":
                active = any(card["type"] == "architecture_evidence" and not card["proof"]["metrics_present"] for card in matched)
            elif rule == "delivery_without_advisory_signal":
                active = any(card["type"] == "delivery_evidence" for card in matched) and not any(card["type"] == "advisory_evidence" for card in evidence_cards)
            elif rule == "build_from_zero_without_scale":
                active = any("scratch" in card["raw_evidence"].lower() or "zero" in card["raw_evidence"].lower() for card in matched) and not any(card["proof"]["scope_present"] for card in matched)
            anti_signals.append(
                {
                    "id": signal["id"],
                    "label": signal["label"],
                    "dimension": signal["dimension"],
                    "status": "active" if active else "clear",
                    "coverage": 100 if active else 0,
                    "confidence": 80 if active else 50,
                    "matched_evidence_ids": [card["id"] for card in matched[:4]],
                }
            )

        return {
            "required_signals": required,
            "supporting_signals": supporting,
            "differentiator_signals": differentiators,
            "anti_signals": anti_signals,
            "not_applicable_signals": [],
        }

    def _score_dimensions(
        self,
        *,
        candidate: dict[str, Any],
        canonical: dict[str, Any],
        signal_snapshot: dict[str, Any],
        merged_pack: dict[str, Any],
        evidence_cards: list[dict[str, Any]],
    ) -> list[dict[str, Any]]:
        dimension_defs = {item["id"]: item for item in merged_pack["dimensions"]}
        evidence_by_dimension: dict[str, list[dict[str, Any]]] = {}
        for group_name in ("required_signals", "supporting_signals", "differentiator_signals"):
            for signal in signal_snapshot[group_name]:
                evidence_by_dimension.setdefault(signal["dimension"], []).append(signal)

        results: list[dict[str, Any]] = []
        for dimension_id, definition in dimension_defs.items():
            signals = evidence_by_dimension.get(dimension_id, [])
            supporting_cards = [card for card in evidence_cards if dimension_id in card.get("dimensions", [])]
            avg_coverage = round(sum(signal["coverage"] for signal in signals) / len(signals)) if signals else 0
            avg_confidence = round(sum(signal["confidence"] for signal in signals) / len(signals)) if signals else 35
            if dimension_id == "cv_clarity_ats":
                structure_bonus = 20 if len(_normalize_lines(canonical["cv_text"])) >= 6 else 0
                section_bonus = 15 if any(marker in canonical["cv_text"].upper() for marker in ["SUMMARY", "EXPERIENCE", "ACHIEVEMENTS", "EDUCATION"]) else 0
                avg_coverage = min(100, avg_coverage + structure_bonus + section_bonus)
            if dimension_id == "role_positioning" and candidate.get("target_role"):
                avg_coverage = min(100, avg_coverage + 15)
            status = (
                "not_applicable"
                if not signals and dimension_id in set(merged_pack.get("not_applicable_defaults", []))
                else ("strong" if avg_coverage >= 75 else "partial" if avg_coverage >= 45 else "weak")
            )
            signals_found = [signal["label"] for signal in signals if signal["coverage"] >= 35][:3]
            signals_missing = [signal["label"] for signal in signals if signal["coverage"] < 35][:3]
            if avg_coverage < 50 and status != "not_applicable" and not signals_missing:
                signals_missing = _dimension_missing_defaults(dimension_id)[:3]
            elif signals_missing:
                signals_missing = _dedupe(signals_missing + _dimension_missing_defaults(dimension_id))[:3]
            why_score_is_not_higher = (
                f"This dimension is still constrained by {', '.join(item.lower() for item in signals_missing[:2])}."
                if signals_missing
                else "The CV points in the right direction, but this dimension still needs sharper proof and clearer packaging."
            )
            results.append(
                {
                    "id": dimension_id,
                    "label": definition["label"],
                    "score": avg_coverage,
                    "weight": definition["weight"],
                    "coverage": avg_coverage,
                    "confidence": avg_confidence,
                    "status": status,
                    "signals_found": signals_found,
                    "signals_missing": signals_missing,
                    "summary": definition["description"],
                    "why_score_is_not_higher": why_score_is_not_higher,
                    "next_best_action": _dimension_next_best_action(dimension_id),
                    "supporting_evidence_ids": [card["id"] for card in supporting_cards[:4]],
                }
            )
        return results

    def _score_primary_scores(
        self,
        *,
        candidate: dict[str, Any],
        canonical: dict[str, Any],
        signal_snapshot: dict[str, Any],
        dimension_states: list[dict[str, Any]],
        support_level: str,
    ) -> dict[str, Any]:
        by_id = {item["id"]: item for item in dimension_states}
        required = signal_snapshot["required_signals"]
        supporting = signal_snapshot["supporting_signals"]
        differentiators = signal_snapshot["differentiator_signals"]
        anti_signals = signal_snapshot["anti_signals"]

        required_avg = round(sum(signal["coverage"] for signal in required) / len(required)) if required else 0
        supporting_avg = round(sum(signal["coverage"] for signal in supporting) / len(supporting)) if supporting else 0
        differentiator_avg = round(sum(signal["coverage"] for signal in differentiators) / len(differentiators)) if differentiators else 0
        anti_penalty = sum(12 for signal in anti_signals if signal["status"] == "active")

        profile_strength = round(
            required_avg * 0.4
            + supporting_avg * 0.2
            + differentiator_avg * 0.15
            + by_id["scope_complexity"]["score"] * 0.15
            + max(0, 100 - anti_penalty * 4) * 0.1
        )
        role_fit = round(
            required_avg * 0.45
            + supporting_avg * 0.2
            + by_id["role_positioning"]["score"] * 0.2
            + by_id["scope_complexity"]["score"] * 0.1
            + max(0, 100 - anti_penalty * 5) * 0.05
        )
        proof_strength = round(
            sum(card["proof"]["metrics_present"] + card["proof"]["scope_present"] + card["proof"]["ownership_present"] + card["proof"]["recency_present"] for card in [
                {
                    "proof": {
                        "metrics_present": by_id["impact_outcomes"]["score"] >= 50,
                        "scope_present": by_id["scope_complexity"]["score"] >= 50,
                        "ownership_present": by_id["evidence_quality"]["score"] >= 50,
                        "recency_present": candidate.get("years_experience", 0) > 0,
                    }
                }
            ]) * 25
        )
        cv_representation_quality = round(
            by_id["cv_clarity_ats"]["score"] * 0.45
            + by_id["role_positioning"]["score"] * 0.25
            + min(100, by_id["evidence_quality"]["score"] + 10) * 0.3
        )
        coverage_pct = round(
            (
                len([signal for signal in required if signal["coverage"] >= 35])
                + len([signal for signal in supporting if signal["coverage"] >= 35])
                + len([signal for signal in differentiators if signal["coverage"] >= 35])
            )
            / max(1, len(required) + len(supporting) + len(differentiators))
            * 100
        )
        if support_level == "curated" and coverage_pct >= 60:
            confidence_label = "High"
            confidence_score = 86
        elif coverage_pct >= 35 and support_level in {"curated", "derived"}:
            confidence_label = "Medium"
            confidence_score = 66
        else:
            confidence_label = "Low"
            confidence_score = 42

        overall_match = round(
            profile_strength * 0.32
            + role_fit * 0.28
            + proof_strength * 0.2
            + cv_representation_quality * 0.2
        )

        return {
            "primary_scores": {
                "profile_strength": max(0, min(100, profile_strength)),
                "role_fit": max(0, min(100, role_fit)),
                "proof_strength": max(0, min(100, proof_strength)),
                "cv_representation_quality": max(0, min(100, cv_representation_quality)),
            },
            "overall_match": max(0, min(100, overall_match)),
            "coverage_pct": coverage_pct,
            "confidence": {"label": confidence_label, "score": confidence_score},
        }

    def _build_gap_map(self, signal_snapshot: dict[str, Any], dimension_states: list[dict[str, Any]]) -> list[dict[str, Any]]:
        gaps: list[dict[str, Any]] = []
        for signal in signal_snapshot["required_signals"] + signal_snapshot["supporting_signals"]:
            if signal["coverage"] >= 70:
                continue
            gaps.append(
                {
                    "id": f"gap::{signal['id']}",
                    "title": signal["label"],
                    "dimension": signal["dimension"],
                    "severity": "high" if signal["tier"] == "required" else "medium",
                    "impact": "high" if signal["tier"] == "required" else "medium",
                    "why_it_matters": signal["expected_evidence"],
                    "evidence_needed": signal["expected_evidence"],
                    "follow_up_hint": signal["question_template"],
                    "benchmark_signal": signal["id"],
                    "question_ids": [],
                }
            )
        for dimension in dimension_states:
            if dimension["score"] >= 55 or dimension["status"] == "not_applicable":
                continue
            gaps.append(
                {
                    "id": f"gap::{dimension['id']}",
                    "title": f"Improve {dimension['label']}",
                    "dimension": dimension["id"],
                    "severity": "high" if dimension["score"] < 40 else "medium",
                    "impact": "high" if dimension["score"] < 40 else "medium",
                    "why_it_matters": dimension["summary"],
                    "evidence_needed": "; ".join(dimension["signals_missing"]) or dimension["label"],
                    "follow_up_hint": dimension["summary"],
                    "benchmark_signal": "",
                    "question_ids": [],
                }
            )
        return gaps[:8]

    def _build_questions(
        self,
        *,
        signal_snapshot: dict[str, Any],
        gap_map: list[dict[str, Any]],
        benchmark_source: dict[str, Any],
        primary_scores: dict[str, Any],
        answers: dict[str, str],
    ) -> list[dict[str, Any]]:
        questions: list[dict[str, Any]] = []
        signals = {}
        for group in ("required_signals", "supporting_signals", "differentiator_signals"):
            for signal in signal_snapshot[group]:
                signals[signal["id"]] = signal

        is_principal_workspace = benchmark_source.get("seniority_pack_id") == "principal" or (
            "principal" in _s(benchmark_source.get("target_role")).lower()
        )
        needs_guided_recovery = (
            primary_scores.get("overall_match", 0) < 70
            or primary_scores.get("coverage_pct", 0) < 70
            or bool(gap_map)
        )

        if is_principal_workspace and needs_guided_recovery:
            principal_questions = [
                {
                    "id": "principal_project_recovery",
                    "title": "Add one principal-level project with architecture depth",
                    "question": "Describe one principal-level project that best proves your architecture depth and technical leadership.",
                    "rationale": "This is usually the fastest way to strengthen technical depth, proof strength, and principal-level credibility.",
                    "why_it_matters": "Principal roles need one clear project that shows architecture depth, trade-offs, scale, and technical ownership.",
                    "placeholder": "Describe one principal-level project with architecture depth, trade-offs, scale, and outcome.",
                    "category": "technical",
                    "status": "pending",
                    "priority": "high",
                    "expected_impact": "high",
                    "dimension": "functional_technical_depth",
                    "question_type": "project_recovery",
                    "benchmark_signal": "",
                    "role_targets": [benchmark_source["target_role"]],
                    "improves_dimensions": ["functional_technical_depth", "evidence_quality"],
                    "estimated_delta": {"global": 6, "roleFit": 7, "proofStrength": 8, "cvRepresentationQuality": 2},
                    "answer_schema": _answer_schema("project_recovery"),
                    "answer_guidance": _answer_guidance("project_recovery"),
                    "example_answer": _answer_example("project_recovery"),
                },
                {
                    "id": "principal_impact_recovery",
                    "title": "Add one measurable transformation result",
                    "question": "What measurable transformation result best proves the impact of your principal-level work?",
                    "rationale": "Without one strong measurable result, the benchmark will keep discounting otherwise strong technical or transformation claims.",
                    "why_it_matters": "Principal-level work still needs a measurable before-and-after result to prove real impact.",
                    "placeholder": "Add one measurable transformation result with baseline, metric, achieved change, timeframe, and your role.",
                    "category": "universal",
                    "status": "pending",
                    "priority": "high",
                    "expected_impact": "high",
                    "dimension": "impact_outcomes",
                    "question_type": "impact_recovery",
                    "benchmark_signal": "",
                    "role_targets": [benchmark_source["target_role"]],
                    "improves_dimensions": ["impact_outcomes", "evidence_quality"],
                    "estimated_delta": {"global": 5, "roleFit": 3, "proofStrength": 8, "cvRepresentationQuality": 2},
                    "answer_schema": _answer_schema("impact_recovery"),
                    "answer_guidance": _answer_guidance("impact_recovery"),
                    "example_answer": _answer_example("impact_recovery"),
                },
                {
                    "id": "principal_target_positioning",
                    "title": "Clarify principal target positioning",
                    "question": "How should this profile position you for principal-level roles in the strongest possible way?",
                    "rationale": "Generic principal positioning weakens both role fit and CV representation quality.",
                    "why_it_matters": "A sharper principal positioning sentence improves benchmark fit, summary quality, and the role variant immediately.",
                    "placeholder": "Clarify the principal track, strongest domain, scope, why principal, and the preferred positioning sentence.",
                    "category": "universal",
                    "status": "pending",
                    "priority": "high",
                    "expected_impact": "high",
                    "dimension": "role_positioning",
                    "question_type": "target_positioning_clarification",
                    "benchmark_signal": "",
                    "role_targets": [benchmark_source["target_role"]],
                    "improves_dimensions": ["role_positioning", "cv_clarity_ats"],
                    "estimated_delta": {"global": 4, "roleFit": 7, "proofStrength": 1, "cvRepresentationQuality": 5},
                    "answer_schema": _answer_schema("target_positioning_clarification"),
                    "answer_guidance": _answer_guidance("target_positioning_clarification"),
                    "example_answer": _answer_example("target_positioning_clarification"),
                },
            ]
            for question in principal_questions:
                if not answers.get(question["id"]):
                    questions.append(question)

        for gap in gap_map:
            signal = signals.get(gap["benchmark_signal"])
            if not signal:
                continue
            if answers.get(signal["id"]):
                continue
            estimated_delta = _delta_template(signal["dimension"], gap["severity"])
            question_type = signal.get("question_family", "universal")
            questions.append(
                {
                    "id": signal["id"],
                    "title": signal["label"],
                    "question": signal["question_template"],
                    "rationale": gap["why_it_matters"],
                    "why_it_matters": gap["why_it_matters"],
                    "placeholder": f"Add evidence for {signal['label'].lower()}",
                    "category": signal["question_family"],
                    "status": "pending",
                    "priority": gap["severity"],
                    "expected_impact": gap["impact"],
                    "dimension": signal["dimension"],
                    "question_type": question_type,
                    "benchmark_signal": signal["id"],
                    "role_targets": [benchmark_source["target_role"]],
                    "improves_dimensions": [signal["dimension"]],
                    "estimated_delta": estimated_delta,
                    "answer_schema": _answer_schema(question_type),
                    "answer_guidance": _answer_guidance(question_type),
                    "example_answer": _answer_example(question_type),
                }
            )

        if not questions and gap_map:
            for gap in gap_map:
                fallback_id = f"question::{gap['id']}"
                if answers.get(fallback_id):
                    continue
                question_family = "technical" if gap["dimension"] == "functional_technical_depth" else "managerial" if gap["dimension"] == "scope_complexity" else "universal"
                estimated_delta = _delta_template(gap["dimension"], gap["severity"])
                questions.append(
                    {
                        "id": fallback_id,
                        "title": gap["title"],
                        "question": gap["follow_up_hint"] or f"What is the clearest proof you can add for {gap['title'].lower()}?",
                        "rationale": gap["why_it_matters"],
                        "why_it_matters": gap["why_it_matters"],
                        "placeholder": f"Add one concrete example that strengthens {gap['title'].lower()}",
                        "category": question_family,
                        "status": "pending",
                        "priority": gap["severity"],
                        "expected_impact": gap["impact"],
                        "dimension": gap["dimension"],
                        "question_type": question_family if gap["dimension"] != "role_positioning" else "target_positioning_clarification",
                        "benchmark_signal": gap["benchmark_signal"],
                        "role_targets": [benchmark_source["target_role"]],
                        "improves_dimensions": [gap["dimension"]],
                        "estimated_delta": estimated_delta,
                        "answer_schema": _answer_schema(question_family if gap["dimension"] != "role_positioning" else "target_positioning_clarification"),
                        "answer_guidance": _answer_guidance(question_family if gap["dimension"] != "role_positioning" else "target_positioning_clarification"),
                        "example_answer": _answer_example(question_family if gap["dimension"] != "role_positioning" else "target_positioning_clarification"),
                    }
                )

        if not benchmark_source["target_role"]:
            questions.insert(
                0,
                {
                    "id": "target_role",
                    "title": "Lock the exact target role",
                    "question": "What exact target role should this benchmark optimize for?",
                    "rationale": "A precise role sharpens the benchmark and all derived questions.",
                    "placeholder": "Example: Director of Data Engineering",
                    "category": "universal",
                    "status": "pending",
                    "priority": "high",
                    "expected_impact": "high",
                    "dimension": "role_positioning",
                    "question_type": "target_positioning_clarification",
                    "benchmark_signal": "",
                    "role_targets": [],
                    "improves_dimensions": ["role_positioning"],
                    "estimated_delta": {"global": 4, "roleFit": 7, "proofStrength": 1, "cvRepresentationQuality": 4},
                    "answer_schema": _answer_schema("target_positioning_clarification"),
                    "answer_guidance": _answer_guidance("target_positioning_clarification"),
                    "example_answer": _answer_example("target_positioning_clarification"),
                }
            )
        if primary_scores["coverage_pct"] < 55 and len(questions) < 3:
            general_gap_dimension = "impact_outcomes" if primary_scores["primary_scores"]["proof_strength"] < 60 else "role_positioning"
            general_id = f"question::general::{general_gap_dimension}"
            if not answers.get(general_id):
                questions.append(
                    {
                        "id": general_id,
                        "title": "Recover one stronger proof point",
                        "question": "What is the single strongest measurable result you can add that best proves you are ready for this role?",
                        "rationale": "The benchmark is still light on approved proof, so one stronger example can lift both proof and role credibility.",
                        "why_it_matters": "A stronger proof point improves confidence, benchmark coverage, and the role-aligned CV variant.",
                        "placeholder": "Describe one strong result with metric, scope, timeframe, and your ownership.",
                        "category": "universal",
                        "status": "pending",
                        "priority": "high",
                        "expected_impact": "high",
                        "dimension": general_gap_dimension,
                        "question_type": "impact_recovery" if general_gap_dimension == "impact_outcomes" else "target_positioning_clarification",
                        "benchmark_signal": "",
                        "role_targets": [benchmark_source["target_role"]],
                        "improves_dimensions": [general_gap_dimension],
                        "estimated_delta": _delta_template(general_gap_dimension, "high"),
                        "answer_schema": _answer_schema("impact_recovery" if general_gap_dimension == "impact_outcomes" else "target_positioning_clarification"),
                        "answer_guidance": _answer_guidance("impact_recovery" if general_gap_dimension == "impact_outcomes" else "target_positioning_clarification"),
                        "example_answer": _answer_example("impact_recovery" if general_gap_dimension == "impact_outcomes" else "target_positioning_clarification"),
                    }
                )
        if not questions and needs_guided_recovery:
            emergency_dimension = gap_map[0]["dimension"] if gap_map else "impact_outcomes"
            questions.append(
                {
                    "id": f"question::emergency::{emergency_dimension}",
                    "title": "Recover one missing proof point",
                    "question": _dimension_next_best_action(emergency_dimension),
                    "rationale": "The benchmark still has open gaps, so Insights should not leave this workspace without at least one guided next step.",
                    "why_it_matters": "A low-coverage workspace still needs at least one guided recovery action.",
                    "placeholder": "Add one concrete example that closes the highest-leverage gap.",
                    "category": "universal",
                    "status": "pending",
                    "priority": "high",
                    "expected_impact": "high",
                    "dimension": emergency_dimension,
                    "question_type": "impact_recovery" if emergency_dimension == "impact_outcomes" else "project_recovery",
                    "benchmark_signal": "",
                    "role_targets": [benchmark_source["target_role"]],
                    "improves_dimensions": [emergency_dimension],
                    "estimated_delta": _delta_template(emergency_dimension, "high"),
                    "answer_schema": _answer_schema("impact_recovery" if emergency_dimension == "impact_outcomes" else "project_recovery"),
                    "answer_guidance": _answer_guidance("impact_recovery" if emergency_dimension == "impact_outcomes" else "project_recovery"),
                    "example_answer": _answer_example("impact_recovery" if emergency_dimension == "impact_outcomes" else "project_recovery"),
                }
            )

        deduped_questions: list[dict[str, Any]] = []
        seen_question_ids: set[str] = set()
        for question in questions:
            if question["id"] in seen_question_ids:
                continue
            seen_question_ids.add(question["id"])
            deduped_questions.append(question)
        return deduped_questions[:3]

    def _build_recommended_profile(
        self,
        candidate: dict[str, Any],
        company: dict[str, Any],
        answers: dict[str, str],
        evidence_cards: list[dict[str, Any]],
        scores: dict[str, Any],
    ) -> dict[str, Any]:
        profile = copy.deepcopy(candidate)
        target_role = answers.get("target_role") or candidate.get("target_role") or company.get("role_title") or candidate.get("current_role")
        focus_evidence = [card["raw_evidence"] for card in evidence_cards if card["strength"] in {"strong", "moderate"}][:5]
        skills = _dedupe(candidate.get("skills", []) + company.get("role_requirements", []))[:12]
        summary_parts = [
            f"{target_role} with {candidate.get('years_experience', 0)}+ years leading data and engineering transformation." if candidate.get("years_experience") else f"{target_role} focused on data and engineering transformation.",
            "Brings measurable leadership and delivery outcomes." if scores["primary_scores"]["proof_strength"] >= 50 else "Needs stronger proof density to fully match the target role.",
        ]
        profile["summary"] = " ".join(part for part in summary_parts if part).strip()
        profile["target_role"] = target_role
        profile["skills"] = skills
        if focus_evidence:
            profile["achievements"] = _dedupe(candidate.get("achievements", []) + focus_evidence)[:5]
        return profile

    def _build_cv_variants(
        self,
        *,
        candidate: dict[str, Any],
        company: dict[str, Any],
        benchmark_source: dict[str, Any],
        recommended_profile: dict[str, Any],
        evidence_cards: list[dict[str, Any]],
        gap_map: list[dict[str, Any]],
    ) -> dict[str, Any]:
        focus_evidence = [card for card in evidence_cards if card["strength"] in {"strong", "moderate"}][:6]
        unresolved_gap_ids = [gap["id"] for gap in gap_map[:3]]

        def build_sections(role_focused: bool) -> list[dict[str, Any]]:
            role_title = benchmark_source["target_role"] or recommended_profile.get("target_role") or company.get("role_title")
            summary = recommended_profile["summary"]
            if role_focused and role_title:
                summary = f"{summary} Positioned specifically for {role_title}."
            items = [card["raw_evidence"] for card in focus_evidence]
            return [
                {"id": "header", "title": "Header", "content": recommended_profile.get("name") or "Candidate", "items": [_headline(role_title or recommended_profile["summary"])]},
                {"id": "summary", "title": "Professional Summary", "content": summary, "items": []},
                {"id": "strengths", "title": "Core Strengths", "content": "", "items": recommended_profile.get("skills", [])[:10]},
                {"id": "impact", "title": "Selected Impact Highlights", "content": "", "items": recommended_profile.get("achievements", [])[:5]},
                {"id": "evidence", "title": "Key Evidence", "content": "", "items": items[:5]},
                {"id": "credentials", "title": "Credentials", "content": recommended_profile.get("education", ""), "items": recommended_profile.get("certifications", []) + recommended_profile.get("languages", [])},
            ]

        def build_variant(variant_id: str, role_focused: bool) -> dict[str, Any]:
            sections = build_sections(role_focused)
            return {
                "variant_id": variant_id,
                "title": "Role Variant" if role_focused else "Master CV",
                "description": f"Focused version aligned to {benchmark_source['target_role']}." if role_focused else "Reusable master version for leadership roles in data and engineering.",
                "source_benchmark_fingerprint": benchmark_source["benchmark_source_fingerprint"],
                "evidence_card_ids_used": [card["id"] for card in focus_evidence],
                "unresolved_gap_ids": unresolved_gap_ids,
                "change_summary": f"Strengthened around {benchmark_source['family']} with curated benchmark rules.",
                "approval_state": "draft",
                "export_state": "not_exported",
                "structured_document_model": sections,
                "rendered_text": self.render_cv_text(sections),
                "sections": sections,
            }

        return {
            "master_cv": build_variant("master_cv", False),
            "role_variant_cv": build_variant("role_variant_cv", True),
        }

    def _build_proposed_changes(self, current_profile: dict[str, Any], recommended_profile: dict[str, Any]) -> list[dict[str, Any]]:
        changes = []
        for field, title, category in (
            ("summary", "Strengthen professional summary", "Narrative"),
            ("target_role", "Clarify target role", "Positioning"),
            ("skills", "Refocus core skills", "Capabilities"),
            ("achievements", "Upgrade impact highlights", "Evidence"),
            ("industry", "Clarify domain positioning", "Domain"),
        ):
            if current_profile.get(field) == recommended_profile.get(field):
                continue
            if field in {"summary", "target_role", "industry"} and not _s(recommended_profile.get(field)):
                continue
            if field in {"skills", "achievements"} and not recommended_profile.get(field):
                continue
            changes.append(
                {
                    "id": f"profile::{field}",
                    "title": title,
                    "category": category,
                    "target": "candidate_profile",
                    "field": field,
                    "reason": "Recommended because it improves benchmark strength and role alignment.",
                    "current_value": current_profile.get(field),
                    "proposed_value": recommended_profile.get(field),
                }
            )
        return changes

    def _build_approved_context_preview(
        self,
        *,
        benchmark_source: dict[str, Any],
        primary_scores: dict[str, Any],
        evidence_cards: list[dict[str, Any]],
        recommended_profile: dict[str, Any],
        selected_evidence_ids: list[str],
    ) -> dict[str, Any]:
        approved_ids = set(selected_evidence_ids)
        chosen_cards = [card for card in evidence_cards if card["id"] in approved_ids] if approved_ids else [card for card in evidence_cards if card["strength"] == "strong"][:4]
        reusable_evidence = [card["raw_evidence"] for card in chosen_cards]
        project_evidence = [card["raw_evidence"] for card in chosen_cards if card["type"] in {"project_evidence", "architecture_evidence", "impact_evidence"}]
        focus_areas = _dedupe([benchmark_source["family"], benchmark_source["archetype"], benchmark_source["seniority"]])
        top_role_signals = _dedupe(
            [signal for card in chosen_cards for signal in card["signal_ids"]]
        )[:4]
        summary = " ".join(
            part
            for part in [
                recommended_profile.get("summary", ""),
                f"Benchmark target: {benchmark_source['target_role']}.",
                f"Support level: {benchmark_source['support_level']}.",
            ]
            if part
        ).strip()
        return {
            "summary": summary,
            "focus_areas": focus_areas,
            "reusable_evidence": reusable_evidence,
            "project_evidence": project_evidence,
            "top_role_signals": top_role_signals,
            "benchmark_headline": f"{benchmark_source['target_role']} benchmark overall {primary_scores['overall_match']}/100",
            "approved_change_titles": [],
            "support_level": benchmark_source["support_level"],
        }

    def _build_cv_health(self, dimension_states: list[dict[str, Any]]) -> str:
        weak = [item["label"].lower() for item in dimension_states if item["score"] < 55][:3]
        if not weak:
            return "CV health is strong for the current benchmark."
        return f"CV health is constrained mainly by {', '.join(weak)}."

    def _build_role_match_summary(self, benchmark_source: dict[str, Any], primary_scores: dict[str, Any]) -> str:
        role = benchmark_source["target_role"] or benchmark_source["family"]
        overall = primary_scores["overall_match"]
        if overall >= 80:
            return f"The profile already looks strong for {role}; the next gains are mostly about sharper proof density and packaging."
        if overall >= 65:
            return f"The profile is directionally credible for {role}, but still needs stronger proof in a few benchmark dimensions."
        return f"The profile has raw material for {role}, but still needs clearer positioning and stronger role-specific evidence."

    def _build_top_strengths(self, signal_snapshot: dict[str, Any], evidence_cards: list[dict[str, Any]]) -> list[str]:
        strengths = [signal["label"] for signal in signal_snapshot["required_signals"] if signal["coverage"] >= 70]
        strengths.extend(signal["label"] for signal in signal_snapshot["supporting_signals"] if signal["coverage"] >= 70)
        if len(strengths) < 3:
            strengths.extend(card["summary"] for card in evidence_cards if card["strength"] == "strong")
        return _dedupe(strengths)[:4]

    def _build_top_gaps(self, gap_map: list[dict[str, Any]]) -> list[str]:
        return _dedupe([gap["title"] for gap in gap_map])[:4]

    def _build_interpretation(
        self,
        *,
        benchmark_source: dict[str, Any],
        primary_scores: dict[str, Any],
        top_strengths: list[str],
        top_gaps: list[str],
        support_level: str,
    ) -> str:
        role = benchmark_source.get("target_role") or benchmark_source.get("family") or "this role"
        if support_level == "unsupported":
            return (
                f"This role is outside the current curated benchmark coverage, so Insights is giving structural CV guidance rather than a high-confidence role benchmark for {role}."
            )
        fit = primary_scores["primary_scores"]["role_fit"]
        proof = primary_scores["primary_scores"]["proof_strength"]
        profile = primary_scores["primary_scores"]["profile_strength"]
        strengths = ", ".join(top_strengths[:2]) if top_strengths else "a few early leadership signals"
        gaps = ", ".join(top_gaps[:2]) if top_gaps else "stronger proof density"
        return (
            f"The profile currently scores {profile}/100 for profile strength and {fit}/100 for role fit against {role}. "
            f"Your clearest strengths are {strengths}, while the biggest gains will come from {gaps}. "
            f"Proof strength is {proof}/100, so the best next move is to recover one or two sharper, measurable examples."
        )

    def _build_improvement_plan(
        self,
        *,
        benchmark_source: dict[str, Any],
        primary_scores: dict[str, Any],
        gap_map: list[dict[str, Any]],
        questions: list[dict[str, Any]],
        proposed_changes: list[dict[str, Any]],
        cv_variants: dict[str, Any],
    ) -> tuple[list[dict[str, Any]], dict[str, Any]]:
        steps: list[dict[str, Any]] = []
        for question in questions:
            steps.append(
                {
                    "step_id": f"question::{question['id']}",
                    "title": question["title"],
                    "type": "question",
                    "why_it_matters": question.get("why_it_matters") or question.get("rationale", ""),
                    "improves_dimensions": question.get("improves_dimensions", []),
                    "estimated_delta": question.get("estimated_delta", _delta_template(question.get("dimension", "evidence_quality"), "medium")),
                    "effort": "low" if question.get("priority") != "high" else "medium",
                    "blocking_dependencies": [],
                    "status": "open",
                }
            )

        if proposed_changes:
            steps.append(
                {
                    "step_id": "rewrite::apply_profile",
                    "title": "Approve the strongest profile rewrite",
                    "type": "apply_rewrite",
                    "why_it_matters": "Applying the best positioning and evidence rewrites improves how the current profile is represented in Prepare.",
                    "improves_dimensions": ["role_positioning", "cv_clarity_ats"],
                    "estimated_delta": {"global": 4, "roleFit": 4, "proofStrength": 1, "cvRepresentationQuality": 5},
                    "effort": "low",
                    "blocking_dependencies": [],
                    "status": "available",
                }
            )

        if cv_variants.get("role_variant_cv"):
            steps.append(
                {
                    "step_id": "variant::regenerate_role",
                    "title": "Regenerate the role-aligned CV variant after approvals",
                    "type": "regenerate_variant",
                    "why_it_matters": "The role variant becomes much stronger once new evidence and approved rewrites are folded in.",
                    "improves_dimensions": ["cv_clarity_ats", "role_positioning"],
                    "estimated_delta": {"global": 3, "roleFit": 3, "proofStrength": 1, "cvRepresentationQuality": 4},
                    "effort": "low",
                    "blocking_dependencies": [step["step_id"] for step in steps[:1]],
                    "status": "available",
                }
            )

        ranked_steps = sorted(
            steps,
            key=lambda step: (
                -sum(step["estimated_delta"].values()),
                0 if step["effort"] == "low" else 1,
            ),
        )
        next_actions = ranked_steps[:3]
        target_score = min(92, primary_scores["overall_match"] + sum(step["estimated_delta"].get("global", 0) for step in next_actions))
        return next_actions, {
            "id": f"plan::{benchmark_source.get('benchmark_source_fingerprint', 'draft')}",
            "role_target": benchmark_source.get("target_role") or benchmark_source.get("family"),
            "current_global_score": primary_scores["overall_match"],
            "target_score": target_score,
            "steps": ranked_steps,
            "open_gap_count": len(gap_map),
        }

    def apply_workspace(
        self,
        *,
        analysis: dict[str, Any],
        approved_change_ids: list[str],
        approved_evidence_ids: list[str],
        targets: list[str],
        variant: str | None,
    ) -> dict[str, Any]:
        current_profile = _normalize_candidate(analysis.get("candidate_profile"))
        workspace = analysis.get("workspace", {})
        updated_profile = copy.deepcopy(current_profile)
        approved_changes = set(approved_change_ids)

        for change in workspace.get("proposed_changes", []):
            if change.get("id") not in approved_changes:
                continue
            if change.get("target") != "candidate_profile":
                continue
            updated_profile[change["field"]] = copy.deepcopy(change.get("proposed_value"))

        cv_text = _s(analysis.get("cv_text"))
        if "cv_text" in targets and variant:
            variant_payload = (workspace.get("cv_variants") or {}).get(variant, {})
            cv_text = _s(variant_payload.get("rendered_text") or cv_text)

        preview = self._build_approved_context_preview(
            benchmark_source=workspace.get("benchmark_source", {}),
            primary_scores={
                "overall_match": workspace.get("overall_match", 0),
            },
            evidence_cards=workspace.get("evidence_cards", []),
            recommended_profile=workspace.get("recommended_profile", updated_profile),
            selected_evidence_ids=approved_evidence_ids,
        )
        preview["approved_change_titles"] = [
            change["title"]
            for change in workspace.get("proposed_changes", [])
            if change.get("id") in approved_changes
        ]

        updated_profile["insights_context_summary"] = preview["summary"]
        updated_profile["insights_focus_areas"] = preview["focus_areas"]
        updated_profile["insights_reusable_evidence"] = preview["reusable_evidence"]

        return {
            "candidate_profile": updated_profile,
            "cv_text": cv_text,
            "applied_change_ids": approved_change_ids,
            "approved_evidence_ids": approved_evidence_ids,
            "variant_applied": variant if "cv_text" in targets else None,
            "approved_context_preview": preview,
            "context_index_status": {
                "saved": False,
                "deleted": {"document_chunks": 0},
                "indexed": {"document_chunks": 0},
            },
        }

    def build_docx_export(self, variant: dict[str, Any], candidate_name: str) -> dict[str, str]:
        from docx import Document

        doc = Document()
        sections = variant.get("structured_document_model", [])
        header = sections[0] if sections else None
        if header:
            doc.add_heading(_s(header.get("content")) or candidate_name or "Candidate", level=0)
            for item in header.get("items", []):
                if _s(item):
                    doc.add_paragraph(_s(item))

        for section in sections[1:]:
            title = _s(section.get("title"))
            if title:
                doc.add_heading(title, level=1)
            content = _s(section.get("content"))
            if content:
                doc.add_paragraph(content)
            for item in section.get("items", []):
                clean = _s(item)
                if clean:
                    doc.add_paragraph(clean, style="List Bullet")

        buffer = BytesIO()
        doc.save(buffer)
        filename = f"{_sanitize_filename(candidate_name)}-{variant.get('variant_id', 'cv')}.docx"
        return {
            "filename": filename,
            "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "content_base64": base64.b64encode(buffer.getvalue()).decode("ascii"),
        }

    def render_cv_text(self, sections: list[dict[str, Any]]) -> str:
        blocks: list[str] = []
        for section in sections:
            title = _s(section.get("title"))
            content = _s(section.get("content"))
            items = [_s(item) for item in section.get("items", []) if _s(item)]
            if title == "Header":
                if content:
                    blocks.append(content)
                if items:
                    blocks.append("\n".join(items))
                continue
            lines = [title.upper()] if title else []
            if content:
                lines.append(content)
            lines.extend(f"- {item}" for item in items)
            if lines:
                blocks.append("\n".join(lines))
        return "\n\n".join(blocks).strip()
